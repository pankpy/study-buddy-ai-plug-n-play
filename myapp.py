import streamlit as st
import google.generativeai as genai
import pandas as pd
from docx import Document
from docx.shared import RGBColor
import os
import time
from datetime import datetime

st.set_page_config(page_title="Question Answer Bot", layout="centered")

st.title("GenAI Q&A (Ask Gemini) 📄")
st.markdown("Use this app to get answers to your questions using Google's Gemini model.")

############ API KEY GUIDE #######################
with st.expander("🔐 How to Get Your Gemini API Key", expanded=False):
    st.markdown("""
    1. Go to [Google AI Studio](https://aistudio.google.com/app/apikey).
    2. Sign in with your Google account.
    3. Click **Create API Key**.
    4. Copy the key and paste it below.
    5. If you refresh (Ctrl+F5) this page, the API key will need to be re-entered here.
    """)

api_key = st.text_input("Paste your Gemini API Key here", type="password")

if api_key:
    genai.configure(api_key=api_key)

############ INPUT OPTIONS ############
st.subheader("📝 Enter Your Questions")

col1, col2 = st.columns(2)

with col1:
    text_questions = st.text_area("Enter up to 5 questions (one per line)", height=150)

with col2:
    # uploaded_file = st.file_uploader("Or upload Excel file (.xlsx). Write questions in column A", type=["xlsx"])
    uploaded_file = st.file_uploader(
        "📤 Or upload Excel file (.xlsx)",
        type=["xlsx"],
        help="Please enter your questions in **Column A** (first column) starting from row 2. Only the first column will be processed."
    )

    st.markdown("""
    > ℹ️ **Note:** Make sure your Excel file has questions in **Column A**, starting from **Row 2**. 
    Avoid adding headers or other data in that column.
    """)

############ Parse Questions ############
questions = []

if text_questions:
    questions += [q.strip() for q in text_questions.strip().split('\n') if q.strip()]

if uploaded_file:
    try:
        df = pd.read_excel(uploaded_file)
        col_names = df.columns.tolist()
        first_col = col_names[0]
        excel_questions = df[first_col].dropna().astype(str).tolist()
        questions += excel_questions
    except Exception as e:
        st.error(f"Could not read Excel file: {e}")

questions = questions[:50]  # Limit to 50 max

############ Main Processing ############
def get_answer(user_query):

    user_query = user_query.strip()
    prompt = f"""
            You are an expert assistant that provides clear, concise, and professional answers to user questions.

            User question is delimited by <<<>>>.
            
            user_query = <<<{user_query}>>>
            
            Follow below steps while answering the user's question.            
            1. Avoid using excessive markdown formatting such as asterisks (**), bold text, or bullet points with symbols like '*'. Instead, use plain language, numbered or clearly separated steps, and short paragraphs.
            2. Respond with well-structured, easy-to-read answers in plain text. 
            3. Do not use emoji, markdown syntax, or decorative characters.
            
            Example style:
            
            Q1: How to learn English?
            
            Learning English can be a rewarding journey. Here are some helpful steps:
            
            1. Clear Goals: Set clear goals. Know why you're learning and what you want to achieve.
            2. Practice: Practice regularly. Spend time listening, speaking, reading, and writing every day.
            3. Daily Use: Use English in your daily life. Watch English shows, label objects, or speak with friends.
            4. Track: Track your progress. Use apps or keep a journal.
            
            Keep your responses practical, clear, and human-friendly.
            """
    system_instruction="""You are a helpful, ethical assistant. Do not answer questions that involve illegal activity, hate speech, violence, personal data, or unethical behavior.If a question is unsafe or inappropriate, politely decline to answer."""

    model = genai.GenerativeModel(
        model_name="gemini-2.5-flash",
        generation_config={"temperature": 0.2},
        system_instruction=system_instruction
    )
    chat = model.start_chat(history=[])
    response = chat.send_message(prompt)
    print("DONE")
    return response.text



def generate_docx(questions):
    doc = Document()
    for i, question in enumerate(questions):
        q_run = doc.add_paragraph().add_run(f"Q{i+1}: {question}")
        q_run.bold = True
        q_run.font.color.rgb = RGBColor(255, 0, 0)

        try:
            answer = get_answer(question)
            st.write(f"Answer {i+1}: Success")
        except Exception as e:
            answer = f"Error getting answer: {e}"

        doc.add_paragraph(f"Answer{i+1}: {answer}")
        doc.add_paragraph()  # Spacer

    filename = f"QA_Output_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    doc.save(filename)
    return filename

############# Run & Display ############
if st.button("🧠 Get Answers"):
    if not api_key:
        st.warning("Please enter your Gemini API key.")
    elif not questions:
        st.warning("Please enter or upload at least one question.")
    else:
        with st.spinner("Generating answers..."):
            output_file = generate_docx(questions)
        with open(output_file, "rb") as f:
            st.success("Done! Click below to download the Word file:")
            st.download_button("📥 Download Answers", data=f, file_name=output_file, mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

