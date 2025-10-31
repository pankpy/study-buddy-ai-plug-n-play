# pip install crewai langchain-google-genai google-generativeai python-docx pandas openpyxl

import streamlit as st
import pandas as pd
from docx import Document
from docx.shared import RGBColor, Inches, Pt
from crewai import Agent, Task, Crew, LLM
import os
from datetime import datetime
import google.generativeai as genai
from crewai import Agent, Task, Crew, Process
import re
from dotenv import load_dotenv

load_dotenv()
api_key = os.getenv("GOOGLE_API_KEY")


def get_answer_with_crewai(user_query, gemini_llm):
    try:
        researcher = Agent(
            role="Research Specialist",
            goal="Analyze questions and identify key concepts",
            backstory="Expert tutor who identifies learning needs.",
            verbose=False,
            allow_delegation=False,
            llm=gemini_llm
        )

        writer = Agent(
            role="Content Writer",
            goal="Create clear, student-friendly explanations with proper formatting",
            backstory="Skilled educator who explains concepts simply with clear section headers.",
            verbose=False,
            allow_delegation=False,
            llm=gemini_llm
        )

        quality_checker = Agent(
            role="Quality Checker",
            goal="Ensure accuracy and completeness",
            backstory="Experienced teacher ensuring study material quality.",
            verbose=False,
            allow_delegation=False,
            llm=gemini_llm
        )

        research_task = Task(
            description=f"Analyze: {user_query}\n\nIdentify key concepts.",
            expected_output="Analysis of concepts",
            agent=researcher
        )

        writing_task = Task(
            description=f"""Create student-friendly answer for: {user_query}

                        Use this format:
                        - Start sections with ## SECTION_NAME (e.g., ## Definition, ## Key Points, ## Example)
                        - Use simple language, examples, numbered steps
                        - Include relevant sections like: Definition, Explanation, Key Points, Steps, Example, Summary
                        - No other markdown formatting
                        
                        Example format:
                        ## Definition
                        [explanation here]
                        
                        ## Key Points
                        1. Point one
                        2. Point two
                        
                        ## Example
                        [example here]""",
            expected_output="Clear, comprehensive answer with sections",
            agent=writer,
            context=[research_task]
        )

        quality_task = Task(
            description="Review answer for clarity, accuracy, format. Ensure student-ready.",
            expected_output="Polished final answer",
            agent=quality_checker,
            context=[writing_task]
        )

        crew = Crew(
            agents=[researcher, writer, quality_checker],
            tasks=[research_task, writing_task, quality_task],
            process=Process.sequential,
            verbose=False
        )

        result = crew.kickoff()
        return result.raw if hasattr(result, 'raw') else str(result)
    except Exception as e:
        return f"Error: {str(e)}"


def get_answer_original(user_query, api_key):
    try:
        prompt = f"""You are an expert tutor. Question: {user_query}

                    Provide clear, comprehensive answer with proper formatting.
                    
                    Use this format:
                    - Start sections with ## SECTION_NAME (e.g., ## Definition, ## Key Points, ## Example)
                    - Use simple language, examples, numbered steps
                    - Include relevant sections like: Definition, Explanation, Key Points, Steps, Example, Summary
                    - No other markdown formatting
                    
                    Example format:
                    ## Definition
                    [explanation here]
                    
                    ## Key Points
                    1. Point one
                    2. Point two
                    
                    ## Example
                    [example here]"""

        model = genai.GenerativeModel(
            model_name="gemini-2.5-flash",
            generation_config={"temperature": 0.1}
        )
        response = model.generate_content(prompt)
        return response.text
    except Exception as e:
        return f"Error: {str(e)}"


def format_answer_in_doc(doc, answer_text):
    """Parse answer text and format sections with colored headers, removing markdown"""
    lines = answer_text.split('\n')

    for line in lines:
        line = line.strip()

        if not line:
            continue

        # Check if line is a section header (## Header or ### Header)
        if line.startswith('###'):
            header_text = line.replace('###', '').strip()
            # Remove ** markers
            header_text = header_text.replace('**', '')
            p = doc.add_paragraph()
            run = p.add_run(header_text)
            run.bold = True
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(76, 132, 234)
        elif line.startswith('##'):
            header_text = line.replace('##', '').strip()
            # Remove ** markers
            header_text = header_text.replace('**', '')
            p = doc.add_paragraph()
            run = p.add_run(header_text)
            run.bold = True
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(76, 132, 234)
        else:
            # Regular text - remove bold markers
            clean_line = line.replace('**', '')
            doc.add_paragraph(clean_line)


def generate_docx(questions, use_crewai, api_key, gemini_llm):
    """Generate docx with error handling that saves progress"""
    doc = Document()
    doc.add_heading('Study Notes - AI Study Buddy', 0)
    doc.add_paragraph(f'Generated: {datetime.now().strftime("%B %d, %Y at %I:%M %p")}')
    doc.add_paragraph()

    successfully_processed = 0
    failed_questions = []

    for i, question in enumerate(questions):
        try:
            # Add question
            q_run = doc.add_paragraph().add_run(f"Q{i + 1}: {question}")
            q_run.bold = True
            q_run.font.color.rgb = RGBColor(255, 0, 0)

            # Generate answer
            try:
                answer = get_answer_with_crewai(question, gemini_llm) if use_crewai else get_answer_original(
                    question, api_key)
                st.write(f"✅ Q{i + 1}/{len(questions)} completed {'(Multi-Agent)' if use_crewai else ''}")

                # Add formatted answer
                doc.add_paragraph("Answer:")
                format_answer_in_doc(doc, answer)

                successfully_processed += 1

            except Exception as e:
                error_msg = f"Error generating answer: {str(e)}"
                doc.add_paragraph(f"Answer: {error_msg}")
                st.error(f"❌ Q{i + 1} failed: {str(e)}")
                failed_questions.append((i + 1, question, str(e)))

            # Add separator
            doc.add_paragraph()
            doc.add_paragraph("─" * 50)
            doc.add_paragraph()

        except Exception as outer_error:
            # If even adding the question fails, log it and continue
            st.error(f"❌ Critical error at Q{i + 1}: {str(outer_error)}")
            failed_questions.append((i + 1, question, f"Critical error: {str(outer_error)}"))
            continue

    # Add summary at the end
    doc.add_paragraph()
    doc.add_heading('Generation Summary', level=1)
    summary_para = doc.add_paragraph()
    summary_run = summary_para.add_run(
        f"Successfully processed: {successfully_processed}/{len(questions)} questions"
    )
    summary_run.font.color.rgb = RGBColor(76, 175, 80) if successfully_processed == len(questions) else RGBColor(255,
                                                                                                                 152, 0)

    if failed_questions:
        doc.add_paragraph()
        doc.add_heading('Failed Questions', level=2)
        for q_num, q_text, error in failed_questions:
            fail_para = doc.add_paragraph()
            fail_run = fail_para.add_run(f"Q{q_num}: {q_text[:50]}... - {error}")
            fail_run.font.color.rgb = RGBColor(244, 67, 54)

    # Save file
    filename = f"Study_Notes_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    doc.save(filename)

    return filename, successfully_processed, len(failed_questions)


# ============================================
# STREAMLIT UI CODE
# ============================================

# --- PAGE CONFIG ---
st.set_page_config(
    page_title="Study Buddy AI",
    layout="wide",
    initial_sidebar_state="collapsed",
    page_icon="📚"
)

# --- THEME FUNCTION ---
def apply_theme(theme: str):
    """Injects CSS styles for the selected theme (red, green, blue)."""
    
    if theme == "red":
        header_bg = "linear-gradient(135deg, #ff9a9e 0%, #ff4b4b 100%)"
        button_bg = "linear-gradient(135deg, #ff7b7b 0%, #ff4b4b 100%)"
        button_hover = "linear-gradient(135deg, #ff5252 0%, #ff0000 100%)"
        info_bg = "#fff5f5"
        border_color = "#ff4b4b"
        shadow_color = "rgba(255, 75, 75, 0.3)"

    elif theme == "green":
        header_bg = "linear-gradient(135deg, #a7f3d0 0%, #10b981 100%)"
        button_bg = "linear-gradient(135deg, #34d399 0%, #10b981 100%)"
        button_hover = "linear-gradient(135deg, #22c55e 0%, #059669 100%)"
        info_bg = "#f0fdf4"
        border_color = "#10b981"
        shadow_color = "rgba(16, 185, 129, 0.3)"

    else:  # blue (your original theme)
        header_bg = "linear-gradient(135deg, #667eea 0%, #764ba2 100%)"
        button_bg = "linear-gradient(135deg, #667eea 0%, #764ba2 100%)"
        button_hover = "linear-gradient(135deg, #5a67d8 0%, #6b46c1 100%)"
        info_bg = "#f0f2f6"
        border_color = "#667eea"
        shadow_color = "rgba(102, 126, 234, 0.3)"
    
    # --- Apply theme ---
    st.markdown(f"""
    <style>
        .main > div {{padding-top: 1rem;}}
        .block-container {{padding-top: 1rem; padding-bottom: 1rem;}}
        h1 {{font-size: 1.8rem !important; margin-bottom: 0.3rem !important;}}
        h2 {{font-size: 1.3rem !important; margin-top: 0.5rem !important;}}
        h3 {{font-size: 1.1rem !important; margin: 0.3rem 0 !important;}}
        .stTextArea textarea {{font-size: 14px !important; min-height: 200px !important;}}

        div[data-testid="stButton"] button {{
            background: {button_bg};
            color: white; font-weight: bold; border-radius: 8px;
            padding: 0.6rem 1.5rem; border: none;
            transition: all 0.2s ease-in-out;
        }}
        div[data-testid="stButton"] button:hover {{
            background: {button_hover};
            transform: scale(1.03);
        }}

        .compact-header {{
            background: {header_bg};
            color: white; padding: 1.2rem; border-radius: 10px;
            text-align: center; margin-bottom: 1rem;
            box-shadow: 0 4px 12px {shadow_color};
        }}

        .info-card {{
            background: {info_bg};
            padding: 0.8rem;
            border-radius: 8px;
            border-left: 4px solid {border_color};
            margin: 0.5rem 0;
        }}

        .stat-badge {{
            background: {button_bg};
            color: white;
            padding: 0.5rem 1rem;
            border-radius: 8px;
            text-align: center;
            font-weight: bold;
        }}
    </style>
    """, unsafe_allow_html=True)


# --- THEME SELECTOR (UI) ---
theme_choice = st.radio(
    "🎨 Choose Theme:",
    ["blue", "green", "red"],
    index=0,
    horizontal=True
)

# --- APPLY SELECTED THEME ---
apply_theme(theme_choice)

# --- HEADER ---
st.markdown("""
<div class="compact-header">
    <h1>📚 Study Buddy AI - Learning Assistant</h1>
    <p style="margin:0; font-size:0.9rem; opacity:0.95;">
        Create Study Notes with AI | Powered by Google Gemini & Multi-Agent AI
    </p>
</div>
""", unsafe_allow_html=True)

left_col, right_col = st.columns([2, 1], gap="medium")

with left_col:
    GEMINI_LLM = None

    try:
        if api_key:
            GEMINI_LLM = LLM(model="gemini/gemini-2.5-flash", temperature=0.1, api_key=api_key)
    except:
        pass

    if not api_key:
        st.warning("Please save gemini API key in .env to continue.")
        st.stop()

    if api_key:
        try:
            genai.configure(api_key=api_key)
            os.environ["GOOGLE_API_KEY"] = api_key
        except Exception as e:
            st.error(f"❌ Error: {e}")
            st.stop()

    st.markdown("### ✍️ Enter Your Study Questions")

    tab1, tab2 = st.tabs(["📝 Type Questions", "📤 Upload Excel"])

    with tab1:
        text_questions = st.text_area(
            "One question per line",
            height=150,
            placeholder="What is photosynthesis?\nExplain Newton's laws of motion\nDescribe the water cycle\nWhat is the structure of an atom?\n\n\nOne question per row",
            label_visibility="collapsed"
        )
        if text_questions:
            q_count = len([q for q in text_questions.strip().split('\n') if q.strip()])
            st.success(f"✅ {q_count} Questions Ready")

    with tab2:
        st.markdown("""
            <div class="info-box">
            <strong>📌 Excel Format:</strong><br>
            • Put all questions in <strong>Column A</strong><br>
            • Start from <strong>Row 2</strong> (Row 1 can be header)<br>
            • One question per row<br>
            """, unsafe_allow_html=True)

        uploaded_file = st.file_uploader(
            "Upload .xlsx file (Questions in Column A, starting Row 2)",
            type=["xlsx"],
            label_visibility="collapsed"
        )
        if uploaded_file:
            st.info("✅ File uploaded successfully")

    st.markdown("### ⚙️ Generation Settings")

    use_crewai = st.checkbox("🤖 Multi-Agent AI", value=True, help="Use 3 AI agents for better answers")

    # Parse Questions
    questions = []
    if text_questions:
        questions += [q.strip() for q in text_questions.strip().split('\n') if q.strip()]

    if uploaded_file:
        try:
            df = pd.read_excel(uploaded_file)
            if len(df.columns) > 0:
                excel_questions = df[df.columns[0]].dropna().astype(str).tolist()
                questions += excel_questions
                st.success(f"✅ Loaded {len(excel_questions)} questions from Excel")
        except Exception as e:
            st.error(f"❌ Excel error: {e}")

    questions = questions[:50]

    if st.button("🚀 Generate My Study Notes Now!", use_container_width=True):
        if not api_key:
            st.warning("⚠️ Please enter your Gemini API key")
        elif not questions:
            st.warning("⚠️ Please enter or upload questions")
        else:
            progress_placeholder = st.empty()

            try:
                with st.spinner("🤖 Generating study notes..."):
                    output_file, success_count, fail_count = generate_docx(questions, use_crewai, api_key, GEMINI_LLM)

                if fail_count == 0:
                    st.balloons()
                    st.success(f"🎉 All {success_count} questions processed successfully!")
                else:
                    st.warning(f"⚠️ Completed with {success_count} successful, {fail_count} failed")

                with open(output_file, "rb") as f:
                    st.download_button(
                        "📥 Download Study Notes (Word Document)",
                        data=f,
                        file_name=output_file,
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        use_container_width=True
                    )

                if fail_count > 0:
                    st.info("💡 Check the end of the document for failed questions summary")
                else:
                    st.info("💡 Tip: Print and highlight key points for better retention!")

            except Exception as e:
                st.error(f"❌ Critical Error: {str(e)}")

with right_col:
    st.markdown("""
    <div class="info-card">
    <strong>✨ Features:</strong><br>
    • AI-generated study notes<br>
    • Excel file support<br>
    • Multi-agent analysis<br>
    • Word document export
    </div>
    """, unsafe_allow_html=True)

    st.markdown("""
    <div class="info-card">
    <strong>📚 Perfect For:</strong><br>
    • Exam/Interview/Assignment Preparation<br>
    • Concept clarification<br>
    • Note making<br>
    </div>
    """, unsafe_allow_html=True)

    st.markdown("""
    <div class="info-card">
    <strong>🎓 All Subjects:</strong><br>
    Science • Math • History<br>
    Literature • Computer Science<br>
    Languages • And more!<br>
    </div>
    """, unsafe_allow_html=True)

    if api_key:
        if st.button("🔍 Test Connection", use_container_width=True):
            try:
                model = genai.GenerativeModel(model_name="gemini-2.5-flash")
                response = model.generate_content("Say 'OK'")
                st.success("✅ Connection OK!")
            except Exception as e:
                st.error(f"❌ Failed: {e}")

st.markdown(
    "<p style='text-align: center; color: #666; font-size: 0.8rem;'>Made with ❤️ for students | Study Smart, Not Hard</p>",
    unsafe_allow_html=True)
